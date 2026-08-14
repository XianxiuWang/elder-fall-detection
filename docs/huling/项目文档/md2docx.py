"""
将 护龄-详细项目计划书.md 转换为 Word (.docx) 文档。
"""
import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
import os

INPUT_MD = r"D:\Users\wangxianxiu\.openclaw\workspace\huling_model\项目文档\护龄-详细项目计划书.md"
OUTPUT_DOCX = r"D:\Users\wangxianxiu\.openclaw\workspace\huling_model\项目文档\护龄-详细项目计划书.docx"

doc = Document()

# =========== 页面设置 ===========
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# =========== 样式定义 ===========
style = doc.styles['Normal']
font = style.font
font.name = '宋体'
font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# 标题样式
for i in range(1, 5):
    heading_style = doc.styles[f'Heading {i}']
    heading_font = heading_style.font
    heading_font.name = '黑体'
    heading_style.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    heading_font.color.rgb = RGBColor(0, 0, 0)
    if i == 1:
        heading_font.size = Pt(22)
    elif i == 2:
        heading_font.size = Pt(16)
    elif i == 3:
        heading_font.size = Pt(13)
    else:
        heading_font.size = Pt(12)

# 创建代码块样式
code_style = doc.styles.add_style('CodeBlock', WD_STYLE_TYPE.PARAGRAPH)
code_font = code_style.font
code_font.name = 'Consolas'
code_font.size = Pt(8.5)
code_font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
code_para = code_style.paragraph_format
code_para.space_before = Pt(3)
code_para.space_after = Pt(3)
code_para.left_indent = Cm(0.8)

def add_code_block(doc, code_text):
    """添加代码块（灰色背景效果）"""
    lines = code_text.strip().split('\n')
    for line in lines:
        p = doc.add_paragraph(line, style='CodeBlock')
        # 添加灰色底纹
        shading_elm = p.paragraph_format.element.get_or_add_pPr()
        shd = shading_elm.makeelement(qn('w:shd'), {
            qn('w:fill'): 'F0F0F0',
            qn('w:val'): 'clear'
        })
        shading_elm.append(shd)

def add_paragraph_with_mixed_format(doc, text):
    """处理行内加粗 `**text**` 和行内代码 `code`"""
    # 清理颜色标签
    text = re.sub(r'\*\*([^*]+)\*\*', r'__BOLD__\1__ENDBOLD__', text)
    text = re.sub(r'`([^`]+)`', r'__CODE__\1__ENDCODE__', text)
    
    # 分割
    parts = re.split(r'(__BOLD__|__ENDBOLD__|__CODE__|__ENDCODE__)', text)
    p = doc.add_paragraph()
    bold_active = False
    code_active = False
    for part in parts:
        if part == '__BOLD__':
            bold_active = True
            continue
        elif part == '__ENDBOLD__':
            bold_active = False
            continue
        elif part == '__CODE__':
            code_active = True
            continue
        elif part == '__ENDCODE__':
            code_active = False
            continue
        run = p.add_run(part)
        if bold_active:
            run.bold = True
        if code_active:
            run.font.name = 'Consolas'
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0xC7, 0x25, 0x4E)
    return p


# =========== 解析 Markdown 并生成 Word ===========
with open(INPUT_MD, 'r', encoding='utf-8') as f:
    lines = f.readlines()

i = 0
in_code_block = False
code_buffer = []
in_table = False
table_buffer = []
in_quote = False

# 状态标记，用于跳过已处理的区域
skip_until_empty = False

while i < len(lines):
    line = lines[i].rstrip()
    
    # ===== 代码块 =====
    if line.startswith('```') and not in_code_block:
        in_code_block = True
        code_buffer = []
        i += 1
        continue
    elif line.startswith('```') and in_code_block:
        in_code_block = False
        if code_buffer:
            add_code_block(doc, '\n'.join(code_buffer))
            doc.add_paragraph()  # 空行
        code_buffer = []
        i += 1
        continue
    elif in_code_block:
        code_buffer.append(line)
        i += 1
        continue
    
    # ===== 水平线 =====
    if line.strip() == '---' or line.strip() == '***':
        doc.add_paragraph('_' * 60)
        i += 1
        continue
    
    # ===== 标题 =====
    heading_match = re.match(r'^(#{1,5})\s+(.+)$', line)
    if heading_match:
        level = len(heading_match.group(1))
        title = heading_match.group(2)
        # 移除行内格式
        title = re.sub(r'\*\*([^*]+)\*\*', r'\1', title)
        title = re.sub(r'`([^`]+)`', r'\1', title)
        doc.add_heading(title, level=level)
        i += 1
        continue
    
    # ===== 表格（检测表头分隔行）=====
    # 如果下一行是 |---|---|... 的模式
    if i + 1 < len(lines) and re.match(r'^\|[\s\-:|]+\|$', lines[i+1].strip()):
        # 表头
        header_cells = [c.strip() for c in line.strip().strip('|').split('|')]
        # 跳过对齐行
        i += 2
        # 收集数据行
        rows = [header_cells]
        while i < len(lines) and lines[i].strip().startswith('|'):
            row_cells = [c.strip() for c in lines[i].strip().strip('|').split('|')]
            rows.append(row_cells)
            i += 1
        
        # 创建表格
        num_cols = len(rows[0])
        table = doc.add_table(rows=len(rows), cols=num_cols, style='Light Grid Accent 1')
        for r_idx, row_data in enumerate(rows):
            for c_idx, cell_text in enumerate(row_data):
                if c_idx < num_cols:
                    cell = table.cell(r_idx, c_idx)
                    cell.text = cell_text
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.size = Pt(9)
                            if r_idx == 0:
                                run.bold = True
        doc.add_paragraph()  # 表后空行
        continue
    
    # ===== 空行 =====
    if line.strip() == '':
        doc.add_paragraph()
        i += 1
        continue
    
    # ===== 普通段落（可能含行内格式）=====
    # 跳过纯链接图片行
    if re.match(r'^!\[.*\]\(.*\)$', line.strip()):
        i += 1
        continue
    # 跳过纯水平线
    if re.match(r'^[\-=\*_]{5,}$', line.strip()):
        i += 1
        continue
    
    # 普通文本段落
    add_paragraph_with_mixed_format(doc, line)
    i += 1

# =========== 保存 ===========
doc.save(OUTPUT_DOCX)
print(f"[OK] Word document saved: {OUTPUT_DOCX}")
print(f"     File size: {os.path.getsize(OUTPUT_DOCX) / 1024:.1f} KB")
